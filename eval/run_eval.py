#!/usr/bin/env python3
"""Round-trip visual accuracy eval for this pdf2docx fork.

Samples are discovered recursively under samples/ -- any directory
(at any nesting depth) containing original.docx is treated as one sample,
named by its path relative to samples/ (e.g. vi/unicode/arial).
Each sample needs:
    original.docx - ground-truth DOCX the conversion should match
    input.pdf     - optional; the PDF to convert. If missing, one is
                     rendered from original.docx via LibreOffice and used
                     as the conversion source instead.

For each sample this script:
    1. renders original.docx -> PDF via LibreOffice (ground truth + input.pdf fallback)
    2. converts that PDF (or the provided input.pdf) -> output.docx using the local pdf2docx fork
    3. renders output.docx -> PDF via LibreOffice
    4. rasterizes each page of both rendered PDFs and scores each pair with SSIM
    5. diffs original.docx vs output.docx text content two ways:
       - strict: paragraph/table-row level (text_diff.txt)
       - loose: whole-document word sequence, whitespace/line-break-insensitive
         (text_diff_loose.txt) -- reflects how much text a user would retype

Results (rendered PDFs, page images, diffs, scores.json) are written to
results/<sample>/ so they can be inspected between iterations.
"""
import difflib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
import numpy as np
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from skimage.metrics import structural_similarity as ssim

from pdf2docx import Converter

EVAL_DIR = Path(__file__).resolve().parent
SAMPLES_DIR = EVAL_DIR / "samples"
RESULTS_DIR = EVAL_DIR / "results"
RENDER_DPI = 150


def convert_pdf_to_docx(pdf_path: Path, docx_path: Path) -> None:
    cv = Converter(str(pdf_path))
    try:
        cv.convert(str(docx_path))
    finally:
        cv.close()


def iter_block_items(doc: Document):
    """Yield paragraphs and tables in true document reading order.

    doc.paragraphs / doc.tables each return their kind in document order but
    as two separate lists, discarding how paragraphs and tables interleave
    (e.g. a caption paragraph followed by a table followed by more prose).
    Walking doc.element.body directly preserves that order.
    """
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def extract_text_lines(docx_path: Path) -> list[str]:
    doc = Document(str(docx_path))
    lines = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                lines.append(text)
        else:
            for row in block.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    lines.append(" | ".join(cells))
    return lines


_LIST_MARKER_GLYPHS = {
    '•', '●', '▪', '■', '♦', '▸', '★',
    '○', '◦', '□', '▹', '☆',
    '‣', '◉', '➤', '➢',
}
_LIST_MARKER_NUMBER_RE = re.compile(r'^\(?\d{1,3}[.\)]$')


def _is_list_marker_token(word: str) -> bool:
    """True if `word` is a leftover list-marker glyph/number (e.g. "•", "1.")
    rather than real content. A real Word list (<w:numPr>) never puts its
    marker into paragraph.text at all -- python-docx/LibreOffice render it
    from the numbering definition -- so a literal marker character sitting in
    extracted text is, at worst, a cosmetic one-character leftover a user
    would delete, not something they'd have to retype. Filtering it out of
    both sides before comparing makes text_sim reflect actual retyping
    burden instead of penalizing this either way.
    """
    stripped = word.replace('​', '').strip()
    return stripped in _LIST_MARKER_GLYPHS or bool(_LIST_MARKER_NUMBER_RE.match(stripped))


def extract_all_words(docx_path: Path) -> list[str]:
    """Whitespace-insensitive word sequence of a docx's entire visible text,
    in document order -- ignores paragraph/line boundaries and table cell
    structure entirely, since neither affects how much text a user would
    have to retype. e.g. "Foo Bar" split across two paragraphs as "Foo" /
    "Bar" collapses to the same word sequence as one paragraph "Foo Bar".
    Leftover list-marker tokens (see `_is_list_marker_token`) are dropped
    from both sides of the comparison, not just here.
    """
    doc = Document(str(docx_path))
    parts = []
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            parts.append(block.text)
        else:
            for row in block.rows:
                for cell in row.cells:
                    parts.append(cell.text)
    words = " ".join(parts).split()
    return [w for w in words if not _is_list_marker_token(w)]


def render_word_diff(orig_words: list[str], out_words: list[str], sm: difflib.SequenceMatcher) -> str:
    """Readable diff of only the real content differences between two word
    sequences, ignoring anything that matches (line breaks/whitespace already
    collapsed away by extract_all_words)."""
    chunks = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        if tag in ("replace", "delete"):
            chunks.append("- " + " ".join(orig_words[i1:i2]))
        if tag in ("replace", "insert"):
            chunks.append("+ " + " ".join(out_words[j1:j2]))
    return "\n".join(chunks)


def text_diff(original_docx: Path, output_docx: Path, diff_path: Path, loose_diff_path: Path) -> dict:
    """Two complementary diffs between two DOCX files' text content. Catches
    issues SSIM can miss the cause of (e.g. Vietnamese diacritics garbled by a
    font/encoding bug) by pointing at the exact text that changed.

    - text_sim_strict: paragraph/table-row level -- sensitive to a paragraph
      being split/merged differently (e.g. a heading reflowed across two
      `<w:p>` instead of one), even when the actual words are identical.
    - text_sim: whole-document, whitespace-normalized word sequence --
      ignores line/paragraph boundaries entirely, only cares whether the same
      words appear in the same order with as few insertions/deletions as
      possible. This is the metric that reflects how much a user would
      actually have to retype.
    """
    orig_lines = extract_text_lines(original_docx)
    out_lines = extract_text_lines(output_docx)

    sm_strict = difflib.SequenceMatcher(a=orig_lines, b=out_lines)
    diff = list(difflib.unified_diff(
        orig_lines, out_lines, fromfile="original.docx", tofile="output.docx", lineterm=""
    ))
    diff_path.write_text("\n".join(diff), encoding="utf-8")
    changed_lines = sum(1 for line in diff if line.startswith(("+", "-"))
                         and not line.startswith(("+++", "---")))

    orig_words = extract_all_words(original_docx)
    out_words = extract_all_words(output_docx)
    sm_loose = difflib.SequenceMatcher(a=orig_words, b=out_words)
    loose_diff_path.write_text(render_word_diff(orig_words, out_words, sm_loose), encoding="utf-8")
    changed_words = sum(max(i2 - i1, j2 - j1) for tag, i1, i2, j1, j2 in sm_loose.get_opcodes()
                         if tag != "equal")

    return {
        "text_sim_strict": sm_strict.ratio(),
        "text_sim": sm_loose.ratio(),
        "n_original_lines": len(orig_lines),
        "n_output_lines": len(out_lines),
        "changed_lines": changed_lines,
        "n_original_words": len(orig_words),
        "n_output_words": len(out_words),
        "changed_words": changed_words,
    }


def docx_to_pdf(docx_path: Path, out_dir: Path) -> Path:
    # Mixed-locale environments (e.g. LC_TIME=vi_VN with LANG=en_US.UTF-8) crash
    # LibreOffice's SvtSysLocaleOptions init, and a shared user profile can
    # corrupt/lock across runs -- force a clean locale and a scratch profile.
    env = {**os.environ, "LC_ALL": "en_US.UTF-8"}
    with tempfile.TemporaryDirectory(prefix="lo_profile_") as profile_dir:
        result = subprocess.run(
            [
                "soffice", "--headless", "--norestore",
                f"-env:UserInstallation=file://{profile_dir}",
                "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path),
            ],
            env=env,
            capture_output=True,
            text=True,
        )
    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice failed converting {docx_path} "
            f"(exit {result.returncode}): {result.stderr}"
        )
    pdf_path = out_dir / (docx_path.stem + ".pdf")
    if not pdf_path.exists():
        raise RuntimeError(f"LibreOffice did not produce {pdf_path}")
    return pdf_path


def render_pages(pdf_path: Path, out_dir: Path, prefix: str) -> list[Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    paths = []
    doc = fitz.open(str(pdf_path))
    zoom = RENDER_DPI / 72
    matrix = fitz.Matrix(zoom, zoom)
    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=matrix)
        path = out_dir / f"{prefix}_p{i + 1:02d}.png"
        pix.save(str(path))
        paths.append(path)
    doc.close()
    return paths


def compare_page(ref_path: Path, out_path: Path, diff_path: Path) -> float:
    import cv2

    ref = cv2.imread(str(ref_path), cv2.IMREAD_GRAYSCALE)
    out = cv2.imread(str(out_path), cv2.IMREAD_GRAYSCALE)

    h = max(ref.shape[0], out.shape[0])
    w = max(ref.shape[1], out.shape[1])
    ref_p = cv2.copyMakeBorder(ref, 0, h - ref.shape[0], 0, w - ref.shape[1],
                                cv2.BORDER_CONSTANT, value=255)
    out_p = cv2.copyMakeBorder(out, 0, h - out.shape[0], 0, w - out.shape[1],
                                cv2.BORDER_CONSTANT, value=255)

    score, diff = ssim(ref_p, out_p, full=True)
    diff_img = (255 - (diff * 255).astype(np.uint8))
    cv2.imwrite(str(diff_path), diff_img)
    return float(score)


def run_sample(name: str) -> dict:
    sample_dir = SAMPLES_DIR / name
    original_docx = sample_dir / "original.docx"
    if not original_docx.exists():
        raise FileNotFoundError(f"Sample '{name}' needs original.docx in {sample_dir}")

    out_dir = RESULTS_DIR / name
    if out_dir.exists():
        shutil.rmtree(out_dir)
    out_dir.mkdir(parents=True)

    print(f"[{name}] rendering original.docx via LibreOffice (ground-truth PDF) ...")
    original_pdf = docx_to_pdf(original_docx, out_dir)

    input_pdf = sample_dir / "input.pdf"
    if not input_pdf.exists():
        # LibreOffice's docx->pdf rendering isn't perfectly reproducible run to
        # run (small text-layout jitter), so a freshly-rendered PDF here would
        # make scores drift between runs for no real reason. Cache it back
        # into the sample dir as a real input.pdf so it's fixed from now on,
        # same as a user-provided one.
        print(f"[{name}] no input.pdf provided -- caching the rendered original.docx as input.pdf")
        shutil.copy(original_pdf, input_pdf)
    conversion_source = input_pdf

    output_docx = out_dir / "output.docx"
    print(f"[{name}] converting -> output.docx ...")
    convert_pdf_to_docx(conversion_source, output_docx)

    print(f"[{name}] diffing text content (original.docx vs output.docx) ...")
    text_result = text_diff(
        original_docx, output_docx, out_dir / "text_diff.txt", out_dir / "text_diff_loose.txt"
    )

    print(f"[{name}] rendering output.docx via LibreOffice ...")
    output_pdf = docx_to_pdf(output_docx, out_dir)

    print(f"[{name}] rasterizing pages ...")
    ref_pages = render_pages(original_pdf, out_dir / "pages", "original")
    out_pages = render_pages(output_pdf, out_dir / "pages", "output")

    n_pages = max(len(ref_pages), len(out_pages))
    page_scores = []
    diff_dir = out_dir / "diffs"
    diff_dir.mkdir(exist_ok=True)
    for i in range(n_pages):
        if i >= len(ref_pages) or i >= len(out_pages):
            page_scores.append(0.0)
            continue
        diff_path = diff_dir / f"diff_p{i + 1:02d}.png"
        score = compare_page(ref_pages[i], out_pages[i], diff_path)
        page_scores.append(score)

    result = {
        "sample": name,
        "n_pages_original": len(ref_pages),
        "n_pages_output": len(out_pages),
        "page_scores": page_scores,
        "mean_score": sum(page_scores) / len(page_scores) if page_scores else 0.0,
        **text_result,
    }
    (out_dir / "scores.json").write_text(json.dumps(result, indent=2, ensure_ascii=False))
    return result


def print_score_table(all_results: list[dict]) -> None:
    cols = [
        ("sample", 30, "s"), ("ssim", 8, ".4f"), ("text_sim", 10, ".4f"),
        ("text_strict", 12, ".4f"), ("changed", 9, "d"), ("pages(o/out)", 13, "s"),
    ]
    header = " | ".join(f"{name:{'<' if fmt == 's' else '>'}{w}}" for name, w, fmt in cols)
    sep = "-+-".join("-" * w for _, w, _ in cols)
    print(header)
    print(sep)
    for r in all_results:
        pages = f"{r['n_pages_original']}/{r['n_pages_output']}"
        print(f"{r['sample']:<30} | {r['mean_score']:>8.4f} | {r['text_sim']:>10.4f} | "
              f"{r['text_sim_strict']:>12.4f} | {r['changed_lines']:>9d} | {pages:>13}")
    print(sep)
    overall_ssim = sum(r["mean_score"] for r in all_results) / len(all_results)
    overall_text = sum(r["text_sim"] for r in all_results) / len(all_results)
    overall_strict = sum(r["text_sim_strict"] for r in all_results) / len(all_results)
    print(f"{'OVERALL':<30} | {overall_ssim:>8.4f} | {overall_text:>10.4f} | "
          f"{overall_strict:>12.4f} | {'':>9} | {'':>13}")
    print(
        "\nssim: 0.0-1.0, higher is better (1.0 = identical) -- visual round-trip similarity. "
        "text_sim: 0.0-1.0, higher is better -- whole-document word-sequence similarity after "
        "collapsing all whitespace/line-breaks (eval/results/<sample>/text_diff_loose.txt); "
        "reflects how much text a user would actually have to retype, ignoring paragraph/line "
        "splits. text_strict (text_sim_strict): 0.0-1.0, higher is better -- paragraph/table-row "
        "level similarity (eval/results/<sample>/text_diff.txt); sensitive to a paragraph being "
        "split/merged differently even when the words are identical. "
        "changed: absolute diff line count (strict), lower is better, NOT normalized by doc "
        "length -- only meaningful when tracking one sample across runs. "
        "pages: original/output page count from rendering each docx via LibreOffice; "
        "a mismatch usually explains a low ssim (pages compared at the wrong index)."
    )


def main() -> None:
    if shutil.which("soffice") is None:
        sys.exit("soffice (LibreOffice) not found on PATH. Install LibreOffice first.")

    names = sorted(
        str(p.parent.relative_to(SAMPLES_DIR)) for p in SAMPLES_DIR.rglob("original.docx")
    )
    if not names:
        sys.exit(f"No samples found in {SAMPLES_DIR}")

    all_results = []
    failed = []
    for name in names:
        try:
            result = run_sample(name)
        except Exception as e:
            print(f"[{name}] FAILED: {e}")
            failed.append(name)
            continue
        all_results.append(result)
        if result["changed_lines"]:
            print(f"[{name}] see eval/results/{name}/text_diff.txt for the exact diff")

    if all_results:
        print("\n=== Score table ===")
        print_score_table(all_results)
    if failed:
        print(f"\nFailed samples (skipped): {', '.join(failed)}")


if __name__ == "__main__":
    main()
